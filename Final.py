import openpyxl
import re
from docx import Document
import mimetypes
import smtplib
import ssl
from email.message import EmailMessage
import os

#Verificar Agregar productos
def CoincidenciaDeDatos(agregarProducto):
    parametro = re.compile(r'\w,\d(\d)?(\d)?(\.\d\d)?,\d(\d)?(\d)?(\.\d\d)?')
    coincidencias = parametro.search(agregarProducto)
    if (coincidencias == None):
        return False
    return True

def coincidenciaClientes(agregarCliente):
    parametro = re.compile(r'\w,\d(\d)?(\d)?(\d)?(\d)?(\d)?(\d)?(\d)?(\d)?(\d)?-\d(\d)?(\d)?(\d)?(\d)?(\d)?(\d)?(\d)?(\d)?(\d)?(\d)?,')
    coincidencias = parametro.search(agregarCliente)
    if (coincidencias == None):
        return False
    return True

def menuPrincipal():
    print('Menu')
    print('a. Productos') 
    print('b. Clientes') 
    print('c. Pedidos') 
    print('d. Informes') 
    print('e. varios') 

def menuA():
    print('Menu')
    print('1. Agregar producto') 
    print('2. Editar producto') 
    print('3. Eliminar producto') 
    print('4. Listar productos') 
    print('5. Enviar cotización por correo') 

def menuB():
    print('Menu')
    print('1. Agregar Cliente') 
    print('2. Editar Cliente') 
    print('3. Eliminar Cliente') 
    print('4. Listar Cliente') 

def menuC():
    print('Menu')
    print('1. Agregar pedido') 
    print('2. Eliminar pedido') 
    print('3. Listar pedido') 

def menuD():
    print('Menu')
    print('1. Total de venta por cliente') 
    print('2. Total de ventas por producto')

def menuE():
    print('Menu')
    print('1. Crear copia de seguridad de datos')

#EL libro a utilizar y las hojas
libro = openpyxl.load_workbook('Inventario.xlsx')
hoja = libro['Productos']
hoja2 = libro['Clientes']
hoja3 = libro['Pedidos']

#Los titulos de cada apartado en productos, Clientes y Pedidos
hoja['A1'].value = "Producto"
hoja['B1'].value = "Precio"
hoja['C1'].value = "Cantidad"

hoja2['A1'].value = "Nombre"
hoja2['B1'].value = "NIT"
hoja2['C1'].value = "Direccion"

hoja3['A1'].value = "Nombre de cliente"
hoja3['B1'].value = "Nombre de producto"
hoja3['C1'].value = "Cantidad de producto"
hoja3['D1'].value = "Valor del pedido"

#variables para el bucle y el diccionario
menu = 0
diccionarioProducto = {}
diccionarioClientes = {}
diccionarioPedidos = {}

while menu == 0:

    totalDeVenta = 0
    totalDePructo = 0
    contador = 2
    contador1 = 0
    productoPedido = ''

    productos = []
    clientes = []
    pedidos = []

    #Sacar los productos ya en existencia del libro
    for row in range(2, hoja.max_row + 1):
        # explora fila por fila
        producto = hoja["A" + str(row)].value
        precio = hoja["B" + str(row)].value
        existencia = hoja["C"+str(row)].value
        if producto != None and  producto != '' :
            diccionarioProducto['Producto'] = producto
            diccionarioProducto['Precio'] = precio
            diccionarioProducto['Cantidad'] = existencia
            productos.append(diccionarioProducto)
        diccionarioProducto = {}
    
    Filas = hoja.min_row + 1
    for valor in productos: 
        hoja['A' + str(Filas)].value = valor['Producto']
        hoja['B' + str(Filas)].value = valor['Precio']
        hoja['C' + str(Filas)].value = valor['Cantidad']
        Filas = Filas + 1
    hoja['A' + str(Filas)].value = ''
    hoja['B' + str(Filas)].value = ''
    hoja['C' + str(Filas)].value = ''

    #Sacar los cliente ya en existencia del libro
    for row in range(2, hoja2.max_row + 1):
        # explora fila por fila
        Nombre = hoja2["A" + str(row)].value
        NIT = hoja2["B" + str(row)].value
        Direccion = hoja2["C"+str(row)].value
        if Nombre != None:
            diccionarioClientes['Nombre'] = Nombre
            diccionarioClientes['NIT'] = NIT
            diccionarioClientes['Direccion'] = Direccion
            clientes.append(diccionarioClientes)
        diccionarioClientes = {}
    
    Filas = hoja2.min_row + 1
    for valor in clientes: 
        hoja2['A' + str(Filas)].value = valor['Nombre']
        hoja2['B' + str(Filas)].value = valor['NIT']
        hoja2['C' + str(Filas)].value = valor['Direccion']
        Filas = Filas + 1
    hoja2['A' + str(Filas)].value = None
    hoja2['B' + str(Filas)].value = ""
    hoja2['C' + str(Filas)].value = ""

    #Sacar los pedidos ya en existencia del libro
    for row in range(2, hoja3.max_row + 1):
        # explora fila por fila
        nombreDeCliente = hoja3["A" + str(row)].value
        nombreDeProducto = hoja3["B" + str(row)].value
        cantidadDePedido = hoja3["C"+str(row)].value
        valorDelPedido = hoja3["D"+str(row)].value
        if nombreDeCliente != None:
            diccionarioPedidos["Nombre de cliente"] = nombreDeCliente
            diccionarioPedidos["Nombre de producto"] = nombreDeProducto
            diccionarioPedidos["Cantidad de producto"] = cantidadDePedido
            diccionarioPedidos["Valor del pedido"] = valorDelPedido
            pedidos.append(diccionarioPedidos)
        diccionarioPedidos = {}
    
    Filas = hoja3.min_row + 1
    for valor in pedidos: 
        hoja3['A' + str(Filas)].value = valor["Nombre de cliente"]
        hoja3['B' + str(Filas)].value = valor["Nombre de producto"]
        hoja3['C' + str(Filas)].value = valor["Cantidad de producto"]
        hoja3['D' + str(Filas)].value = valor["Valor del pedido"]
        Filas = Filas + 1
    hoja3['A' + str(Filas)].value = None
    hoja3['B' + str(Filas)].value = ""
    hoja3['C' + str(Filas)].value = ""
    hoja3['D' + str(Filas)].value = ""

    #guardar cambios
    libro.save("Inventario.xlsx")

    menuPrincipal()
    eleccion = input('\n')
    
    if eleccion == 'a':

        menuA()
        eleccion2 = input('\n')

        #Agregar producto
        if eleccion2 == '1':
            agregarProducto = input("Ingrese los datos en este orden Producto,Precio,Cantidad\n")
            if CoincidenciaDeDatos(agregarProducto) == True:
                print("Se añadieron los datos correctamente")
            if CoincidenciaDeDatos(agregarProducto) == False:
                print("Los datos no cumplen con los parametros establecidos")

            delimitador = "," 
            agregarProductoSeparado = agregarProducto.split(delimitador)

            #añadir al diccionario
            diccionarioProducto['Producto'] = agregarProductoSeparado[0]
            diccionarioProducto['Precio'] = float(agregarProductoSeparado[1])
            diccionarioProducto['Cantidad'] = float(agregarProductoSeparado[2])
            productos.append(diccionarioProducto)
            diccionarioProducto = {}

            #añadir al libro de excel
            Filas = hoja.min_row + 1
            for valor in productos: 
                hoja['A' + str(Filas)].value = valor['Producto']
                hoja['B' + str(Filas)].value = valor['Precio']
                hoja['C' + str(Filas)].value = valor['Cantidad']
                Filas = Filas + 1
            #guardar cambios
            libro.save("Inventario.xlsx")

        #Editar producto
        if eleccion2 == '2':
            for valor2 in productos:
                print(valor2['Producto'])
            editarProducto = input("Ingrese el producto a editar\n")

            for valor2 in productos:
                if editarProducto == valor2['Producto']:
                    agregarProducto = input("Ingrese los nuevos datos (Producto,Precio,Existencia)\n")
                    if CoincidenciaDeDatos(agregarProducto) == True:
                        print("Se añadieron los datos correctamente")
                    if CoincidenciaDeDatos(agregarProducto) == False:
                        print("Los datos no cumplen con los parametros establecidos")
                    
                    delimitador = "," 
                    agregarProductoSeparado = agregarProducto.split(delimitador)

                    for valor in agregarProductoSeparado : 
                        hoja['A' + str(contador)].value = agregarProductoSeparado[0]
                        hoja['B' + str(contador)].value = float(agregarProductoSeparado[1])
                        hoja['C' + str(contador)].value = float(agregarProductoSeparado[2])

                    #guardar cambios
                    libro.save("Inventario.xlsx")
     
                else:
                    contador = contador + 1
        
        #Eliminar producto
        if eleccion2 == '3':
            for valor2 in productos:
                print(valor2['Producto'])
            eliminarProducto = input("Ingrese el producto a eliminar\n")

            for valor2 in productos:
                if eliminarProducto == valor2['Producto']:
 
                    hoja['A' + str(contador)].value = None
                    hoja['B' + str(contador)].value = None
                    hoja['C' + str(contador)].value = None
                    print('El producto se a eliminado correctamente')
                    contador1 = 1

                    #guardar cambios
                    libro.save("Inventario.xlsx")
     
                else:
                    contador = contador + 1
            if contador1 != 1:
                print('El producto que desea eliminar no existe')

        #Listar producto producto
        if eleccion2 == '4':

            for item in productos:
                print ("Producto: " + item['Producto'])
                print ("Precio: " + str(item['Precio']))
                print ("Cantidad: " + str(item['Cantidad'])+'\n')

        #Enviar cotización por correo
        if eleccion2 == '5':
            for valor2 in clientes:
                print (valor2['Nombre'])
            NombreClienteCotizacion = input('Ingrese el nombre del cliente\n')
            correoClienteCotizacion = input('Ingrese su correo electronico\n')
            for valor3 in productos:
                print (valor3['Producto'])
            productoElegidoCotizacion = input('Ingrese el producto elegido\n')

            for valor2 in clientes:

                for valor3 in productos:
                    if productoElegidoCotizacion == valor3['Producto']:
                        contador = 10
                        precioPedido = valor3['Precio']
                        productoPedido = valor3['Producto']

                if NombreClienteCotizacion == valor2['Nombre'] and  contador == 10:
                    cotizacion = Document()
                    cotizacion.add_heading('Cotización', 0)
                    p = cotizacion.add_paragraph('Estimado '+valor2['Nombre'])
                    p = cotizacion.add_paragraph('El precio de nuestro producto: '+str(productoPedido))
                    p = cotizacion.add_paragraph('es: Q '+str(precioPedido))
                    cotizacion.save('cotizacion.docx')
                    contador1 = 1

            if contador1 != 1:
                print('El producto o cliente que ingreso no existe')
            else:
                #envio por correo
                DIRECCION_DEL_SERVIDOR = "smtp.gmail.com"
                PUERTO = 587
                DIRECCION_DE_ORIGEN = "pruebasalgoritmos123@gmail.com"
                CONTRASENA = 'prueba123'

                #Contenido del mensaje
                mensaje = EmailMessage()
                mensaje["Subject"] = "Cotizacion"
                mensaje["From"] = DIRECCION_DE_ORIGEN
                mensaje["To"] = correoClienteCotizacion

                mensaje.add_alternative(""" 
                <p> 
                <h1>No responder este mensaje</h1>
                </p>
                """, subtype = "html")

                nombre_de_archivo = "cotizacion.docx"
                ctype, encoding = mimetypes.guess_type(nombre_de_archivo)

                if ctype is None or encoding is not None:
                    ctype = 'application/octet-stream'

                tipoPrincipal, subTipo = ctype.split('/', 1)

                with open(nombre_de_archivo, 'rb') as archivoLeido:
                    mensaje.add_attachment(archivoLeido.read(), maintype=tipoPrincipal, subtype = subTipo, filename = nombre_de_archivo)

                context = ssl.create_default_context()

                smtp = smtplib.SMTP(DIRECCION_DEL_SERVIDOR, PUERTO)
                smtp.starttls()
                smtp.login(DIRECCION_DE_ORIGEN, CONTRASENA)
                smtp.send_message(mensaje)

                print('La cotizacin se a enviado correctamente')
            
    if eleccion == 'b':

        menuB()
        eleccion2 = input('\n')

        #Agregar Cliente
        if eleccion2 == '1':
            agregarCliente= input("Ingrese los datos en este orden Nombre y apellido,NIT,Direccion\n")
            if coincidenciaClientes(agregarCliente) == True:
                print("Se añadieron los datos correctamente")
            if coincidenciaClientes(agregarCliente) == False:
                print("Los datos no cumplen con los parametros establecidos")

            delimitador = "," 
            agregarClienteSeparado = agregarCliente.split(delimitador)

            #añadir al diccionario
            diccionarioClientes['Nombre'] = agregarClienteSeparado[0]
            diccionarioClientes['NIT'] = agregarClienteSeparado[1]
            diccionarioClientes['Direccion'] = agregarClienteSeparado[2]
            clientes.append(diccionarioClientes)
            diccionarioClientes = {}

            #añadir al libro de excel
            Filas = hoja2.min_row + 1
            for valor in clientes: 
                hoja2['A' + str(Filas)].value = valor['Nombre']
                hoja2['B' + str(Filas)].value = valor['NIT']
                hoja2['C' + str(Filas)].value = valor['Direccion']
                Filas = Filas + 1

            #guardar cambios
            libro.save("Inventario.xlsx")
    
        #Editar Cliente
        if eleccion2 == '2':
            for valor2 in clientes:
                print(valor2['Nombre'])
            editarCliente = input("Ingrese el cliente a editar\n")

            for valor2 in clientes:
                if editarCliente == valor2['Nombre']:
                    agregarCliente = input("Ingrese los nuevos datos (Nombre y apellido,NIT,Direccion)\n")
                    if coincidenciaClientes(agregarCliente) == True:
                        print("Se añadieron los datos correctamente")
                    if coincidenciaClientes(agregarCliente) == False:
                        print("Los datos no cumplen con los parametros establecidos")
                    
                    delimitador = "," 
                    agregarClienteSeparado = agregarCliente.split(delimitador)

                    for valor in agregarClienteSeparado : 
                        hoja2['A' + str(contador)].value = agregarClienteSeparado[0]
                        hoja2['B' + str(contador)].value = agregarClienteSeparado[1]
                        hoja2['C' + str(contador)].value = agregarClienteSeparado[2]

                    #guardar cambios
                    libro.save("Inventario.xlsx")
     
                else:
                    contador = contador + 1
        
        #Eliminar Clientes
        if eleccion2 == '3':
            for valor2 in clientes:
                print(valor2['Nombre'])
            eliminarCliente = input("Ingrese el nombre del cliente a eliminar\n")

            for valor2 in clientes:
                if eliminarCliente == valor2['Nombre']:
 
                    hoja2['A' + str(contador)].value = None
                    hoja2['B' + str(contador)].value = None
                    hoja2['C' + str(contador)].value = None
                    print('El cliente se a eliminado correctamente')
                    contador1 = 1

                    #guardar cambios
                    libro.save("Inventario.xlsx")
                else:
                    contador = contador + 1

            if contador1 != 1:
                print('El cliente que ingreso no existe')
        
        #Listar Clientes
        if eleccion2 == '4':

            for item in clientes:
                print ("Cliente: " + item['Nombre'])
                print ("NIT: " + str(item['NIT']))
                print ("Direccion: " + str(item['Direccion'])+'\n')

    if eleccion == 'c':

        menuC()
        eleccion2 = input('\n')

        #Agregar pedido
        if eleccion2 == '1':

            #mostrar inventario
            print('Elija el nombre del cliente')
            for valor2 in clientes:
                print (valor2['Nombre'])
            clientePedido = input()
            print('Elija un producto de nuestro inventario')
            for valor3 in productos:
                print (valor3['Producto'])
            productoPedido = input()
            print('Ingrese la cantidad de producto que desea')
            cantidadPedido = input()
            

            for valor2 in clientes:

                for valor3 in productos:
                    if productoPedido == valor3['Producto']:
                        contador = 10
                        precioPedido = valor3['Precio']

                if clientePedido == valor2['Nombre'] and  contador == 10:
                    diccionarioPedidos["Nombre de cliente"] = clientePedido
                    diccionarioPedidos["Nombre de producto"] = productoPedido
                    diccionarioPedidos["Cantidad de producto"] = int(cantidadPedido)
                    diccionarioPedidos["Valor del pedido"] = precioPedido*float(cantidadPedido)
                    pedidos.append(diccionarioPedidos)
                    diccionarioPedidos = {}

                    Filas = hoja.min_row + 1
                    for valor in pedidos: 
                        hoja3['A' + str(Filas)].value = valor["Nombre de cliente"]
                        hoja3['B' + str(Filas)].value = valor["Nombre de producto"]
                        hoja3['C' + str(Filas)].value = valor["Cantidad de producto"]
                        hoja3['D' + str(Filas)].value = valor["Valor del pedido"]
                        Filas = Filas + 1
                        print('El pedido se agrego correctamente')
                        contador1 = 1

                        #guardar cambios
                        libro.save("Inventario.xlsx")

            if contador1 != 1:
                print('El producto o cliente que ingreso no existe')

        #Eliminar pedido
        if eleccion2 == '2':
            for valor2 in pedidos:
                print(valor2['Nombre de cliente'])
            eliminarPedido = input("Ingrese el Nombre de cliente al que corresponde el pedido a eliminar\n")

            for valor2 in pedidos:
                if eliminarPedido == valor2['Nombre de cliente']:
 
                    hoja3['A' + str(contador)].value = None
                    hoja3['B' + str(contador)].value = None
                    hoja3['C' + str(contador)].value = None
                    hoja3['D' + str(contador)].value = None
                    contador1 = 1
                    print('El pedido se elimino correctamente')

                    #guardar cambios
                    libro.save("Inventario.xlsx")
     
                else:
                    contador = contador + 1
            if contador1 != 1:
                print('El cliente que ingreso no existe')
        
        #Listar pedido
        if eleccion2 == '3':

            for item in pedidos:
                print ("Nombre de cliente: " + item['Nombre de cliente'])
                print ("Nombre de producto: " + str(item['Nombre de producto']))
                print ("Cantidad de producto: " + str(item['Cantidad de producto']))
                print ("Valor del pedido: " + str(item['Valor del pedido'])+'\n')
        
    if eleccion == 'd':
        
        menuD()
        eleccion2 = input('\n')

        #Total de venta por cliente
        if eleccion2 == '1':
            for valor2 in pedidos:
                print(valor2['Nombre de cliente'])
            buscarClientes = input("Ingrese el nombre del cliente del que desea saber el total de venta\n")

            for valor2 in pedidos:
                if buscarClientes == valor2['Nombre de cliente']:
                    totalDeVenta = totalDeVenta + valor2['Valor del pedido']
                    contador1 = contador + 1

            print ("Nombre del cliente: " + buscarClientes)
            print ("Cantidad de pedidos: " + str(contador))
            print ("Total de ventas: " + str(totalDeVenta)+'\n')

            if contador1 != 1:
                print('El cliente que ingreso no existe')

        #Total de venta por producto
        if eleccion2 == '2':
            for valor2 in pedidos:
                print(valor2['Nombre de producto'])
            buscarProducto = input("Ingrese el nombre del producto del que desea saber el total de venta\n")

            for valor2 in pedidos:
                if buscarProducto == valor2['Nombre de producto']:
                    totalDeVenta = totalDeVenta + valor2['Valor del pedido']
                    totalDePructo = totalDePructo + valor2['Cantidad de producto']
                    contador1 = 1

            print ("Nombre del producto: " + buscarProducto)
            print ("Cantidad de producto: " + str(totalDePructo))
            print ("Total de ventas: " + str(totalDeVenta)+'\n')

            if contador1 != 1:
                print('El producto que ingreso no existe')

    if eleccion == 'e':

        menuE()
        eleccion2 = input('\n')

        #Crear copia de seguridad de datos
        if eleccion2 == '1':
            libro.save("Inventario_copia.xlsx")

            DIRECCION_DEL_SERVIDOR = "smtp.gmail.com"
            PUERTO = 587
            DIRECCION_DE_ORIGEN = "pruebasalgoritmos123@gmail.com"
            CONTRASENA = 'prueba123'

            #Contenido del mensaje
            mensaje = EmailMessage()
            mensaje["Subject"] = "Copia de seguridad"
            mensaje["From"] = DIRECCION_DE_ORIGEN
            mensaje["To"] = 'jsincala@miumg.edu.gt'

            mensaje.add_alternative(""" 
            <p> 
            <h1>No responder este mensaje</h1>
            </p>
            """, subtype = "html")

            nombre_de_archivo = "Inventario_copia.xlsx"
            ctype, encoding = mimetypes.guess_type(nombre_de_archivo)

            if ctype is None or encoding is not None:
                ctype = 'application/octet-stream'

            tipoPrincipal, subTipo = ctype.split('/', 1)

            with open(nombre_de_archivo, 'rb') as archivoLeido:
                mensaje.add_attachment(archivoLeido.read(), maintype=tipoPrincipal, subtype = subTipo, filename = nombre_de_archivo)

            context = ssl.create_default_context()

            smtp = smtplib.SMTP(DIRECCION_DEL_SERVIDOR, PUERTO)
            smtp.starttls()
            smtp.login(DIRECCION_DE_ORIGEN, CONTRASENA)
            smtp.send_message(mensaje)

            print('La copia de seguridad ha sido creada\n')